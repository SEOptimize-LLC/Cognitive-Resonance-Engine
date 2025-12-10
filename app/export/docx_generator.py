"""
DOCX Report Generator
=====================
Generates comprehensive Word documents from research results.
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.models.research_models import ResearchResults


def generate_docx(results: ResearchResults) -> bytes:
    """
    Generate a comprehensive DOCX report from research results.
    
    Args:
        results: The complete research results
        
    Returns:
        Bytes of the DOCX file
    """
    doc = Document()
    
    # Set up styles
    setup_styles(doc)
    
    # Title
    title = doc.add_heading(
        f"Audience Research Report: {results.client_input.client_name}",
        0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.italic = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Table of Contents Header
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Company Profile",
        "3. Ideal Customer Profiles",
        "4. Value Propositions",
        "5. Pain Point Analysis",
        "6. Customer Journey Maps",
        "7. Cost Summary"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style = 'List Bullet'
    
    doc.add_page_break()
    
    # Executive Summary
    add_executive_summary(doc, results)
    
    # Company Profile
    add_company_profile(doc, results)
    
    # ICPs
    add_icps_section(doc, results)
    
    # Value Propositions
    add_vp_section(doc, results)
    
    # Pain Points
    add_pain_section(doc, results)
    
    # Journey Maps
    add_journey_section(doc, results)
    
    # Cost Summary
    add_cost_section(doc, results)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


def setup_styles(doc: Document):
    """Configure document styles."""
    styles = doc.styles
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 style  
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 76, 153)
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0, 102, 204)


def add_executive_summary(doc: Document, results: ResearchResults):
    """Add executive summary section."""
    doc.add_heading("1. Executive Summary", level=1)
    
    client = results.client_input
    profile = results.company_profile
    
    p = doc.add_paragraph()
    p.add_run("This report presents comprehensive audience research for ")
    p.add_run(f"{client.client_name}").bold = True
    p.add_run(
        f", a {client.business_model.value} company "
        f"in the {client.industry} industry."
    )
    
    if profile and profile.overview:
        doc.add_paragraph(profile.overview)
    
    # Key findings
    doc.add_heading("Key Findings", level=2)
    
    num_icps = len(results.icps)
    doc.add_paragraph(
        f"{num_icps} Ideal Customer Profiles identified",
        style='List Bullet'
    )
    
    high_priority = sum(
        1 for icp_result in results.icps
        if icp_result.icp.segment_priority == "high"
    )
    if high_priority:
        doc.add_paragraph(
            f"{high_priority} High-Priority Segments for targeting",
            style='List Bullet'
        )
    
    # List ICPs
    if results.icps:
        doc.add_heading("Identified ICPs", level=2)
        for icp_result in results.icps:
            icp = icp_result.icp
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{icp.icp_name}").bold = True
            p.add_run(f" ({icp.segment_priority}): {icp.one_liner}")
    
    doc.add_page_break()


def add_company_profile(doc: Document, results: ResearchResults):
    """Add company profile section."""
    doc.add_heading("2. Company Profile", level=1)
    
    profile = results.company_profile
    if not profile:
        doc.add_paragraph("Company profile not available.")
        doc.add_page_break()
        return
    
    # Company name
    doc.add_heading(profile.name, level=2)
    
    # Basic info table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    cells = [
        ("Website", profile.website_url),
        ("Industry", profile.industry),
        ("Business Model", profile.business_model.value)
    ]
    
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = value or "N/A"
    
    doc.add_paragraph()
    
    if profile.overview:
        doc.add_paragraph(profile.overview)
    
    # Products & Services
    if profile.products_services:
        doc.add_heading("Products & Services", level=2)
        for ps in profile.products_services:
            doc.add_heading(ps.name, level=3)
            doc.add_paragraph(ps.description)
            
            if ps.features:
                p = doc.add_paragraph()
                p.add_run("Features: ").bold = True
                p.add_run(", ".join(ps.features))
    
    # Value Propositions
    if profile.stated_value_propositions:
        doc.add_heading("Stated Value Propositions", level=2)
        for vp in profile.stated_value_propositions:
            doc.add_paragraph(vp, style='List Bullet')
    
    # Competitors
    if profile.competitors:
        doc.add_heading("Competitive Landscape", level=2)
        for comp in profile.competitors[:5]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{comp.name}").bold = True
            if comp.description:
                p.add_run(f" - {comp.description}")
    
    doc.add_page_break()


def add_icps_section(doc: Document, results: ResearchResults):
    """Add ICPs section."""
    doc.add_heading("3. Ideal Customer Profiles", level=1)
    
    if not results.icps:
        doc.add_paragraph("No ICPs generated.")
        doc.add_page_break()
        return
    
    for i, icp_result in enumerate(results.icps, 1):
        icp = icp_result.icp
        
        doc.add_heading(f"ICP {i}: {icp.icp_name}", level=2)
        
        # Priority and one-liner
        p = doc.add_paragraph()
        p.add_run("Priority: ").bold = True
        p.add_run(icp.segment_priority.upper())
        
        p = doc.add_paragraph()
        p.add_run(icp.one_liner).italic = True
        
        # Demographics
        doc.add_heading("Demographics/Firmographics", level=3)
        demo = icp.demographics
        
        demo_items = []
        if demo.company_size:
            demo_items.append(f"Company Size: {demo.company_size}")
        if demo.job_titles:
            demo_items.append(f"Job Titles: {', '.join(demo.job_titles)}")
        if demo.industry_verticals:
            demo_items.append(
                f"Industries: {', '.join(demo.industry_verticals)}"
            )
        if demo.geographic_focus:
            demo_items.append(f"Geography: {demo.geographic_focus}")
        if demo.budget_range:
            demo_items.append(f"Budget: {demo.budget_range}")
        
        for item in demo_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Psychographics
        doc.add_heading("Psychographics", level=3)
        psycho = icp.psychographics
        
        psycho_items = [
            f"Decision Style: {psycho.decision_style.value}",
            f"Risk Tolerance: {psycho.risk_tolerance.value}"
        ]
        if psycho.core_values:
            psycho_items.append(f"Core Values: {', '.join(psycho.core_values)}")
        if psycho.fears:
            psycho_items.append(f"Key Fears: {', '.join(psycho.fears)}")
        
        for item in psycho_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Motivations
        if icp.motivations:
            doc.add_heading("Key Motivations", level=3)
            for m in icp.motivations[:5]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(m.statement)
                p.add_run(f" (Intensity: {m.intensity}/10)").italic = True
        
        # Pain Points
        if icp.pain_points:
            doc.add_heading("Key Pain Points", level=3)
            for pp in icp.pain_points[:5]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(pp.statement)
                p.add_run(f" (Severity: {pp.severity}/10)").italic = True
        
        # Narrative
        if icp.detailed_narrative:
            doc.add_heading("Detailed Profile", level=3)
            doc.add_paragraph(icp.detailed_narrative)
        
        if i < len(results.icps):
            doc.add_paragraph()
    
    doc.add_page_break()


def add_vp_section(doc: Document, results: ResearchResults):
    """Add value propositions section."""
    doc.add_heading("4. Value Propositions", level=1)
    
    has_vp = any(r.value_proposition for r in results.icps)
    if not has_vp:
        doc.add_paragraph("Value propositions not analyzed.")
        doc.add_page_break()
        return
    
    for icp_result in results.icps:
        if not icp_result.value_proposition:
            continue
        
        vp = icp_result.value_proposition
        icp = icp_result.icp
        
        doc.add_heading(f"Value Proposition for {icp.icp_name}", level=2)
        
        if vp.value_proposition_statement:
            p = doc.add_paragraph()
            p.add_run(vp.value_proposition_statement).italic = True
        
        if vp.fit_score:
            p = doc.add_paragraph()
            p.add_run("Fit Score: ").bold = True
            p.add_run(f"{vp.fit_score}/100")
        
        # Pain Relievers
        if vp.pain_relievers:
            doc.add_heading("Pain Relievers", level=3)
            for pr in vp.pain_relievers[:5]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{pr.pain_addressed}").bold = True
                p.add_run(f"\n→ {pr.how_relieved}")
        
        # Gain Creators
        if vp.gain_creators:
            doc.add_heading("Gain Creators", level=3)
            for gc in vp.gain_creators[:5]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{gc.gain_created}").bold = True
                p.add_run(f"\n→ {gc.how_created}")
        
        # Differentiators
        if vp.unique_differentiators:
            doc.add_heading("Unique Differentiators", level=3)
            for d in vp.unique_differentiators:
                doc.add_paragraph(d, style='List Bullet')
    
    doc.add_page_break()


def add_pain_section(doc: Document, results: ResearchResults):
    """Add pain points section."""
    doc.add_heading("5. Pain Point Analysis", level=1)
    
    has_pain = any(r.pain_taxonomy for r in results.icps)
    if not has_pain:
        doc.add_paragraph("Pain taxonomy not analyzed.")
        doc.add_page_break()
        return
    
    for icp_result in results.icps:
        if not icp_result.pain_taxonomy:
            continue
        
        pt = icp_result.pain_taxonomy
        icp = icp_result.icp
        
        doc.add_heading(f"Pain Taxonomy: {icp.icp_name}", level=2)
        
        # Functional Pains Table
        if pt.functional_pains:
            doc.add_heading("Functional Pain Points", level=3)
            table = doc.add_table(
                rows=len(pt.functional_pains[:7]) + 1,
                cols=3
            )
            table.style = 'Table Grid'
            
            # Header
            headers = ["Pain Point", "Category", "Severity"]
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
                table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            
            # Data
            for row_idx, p in enumerate(pt.functional_pains[:7], 1):
                table.rows[row_idx].cells[0].text = p.statement
                table.rows[row_idx].cells[1].text = p.category
                table.rows[row_idx].cells[2].text = f"{p.severity}/10"
            
            doc.add_paragraph()
        
        # Forces Analysis
        if pt.forces_analysis:
            fa = pt.forces_analysis
            doc.add_heading("Forces of Progress Analysis", level=3)
            
            if fa.push_factors:
                p = doc.add_paragraph()
                p.add_run("Push Factors: ").bold = True
                p.add_run(", ".join(fa.push_factors[:3]))
            
            if fa.pull_factors:
                p = doc.add_paragraph()
                p.add_run("Pull Factors: ").bold = True
                p.add_run(", ".join(fa.pull_factors[:3]))
            
            if fa.habit_factors:
                p = doc.add_paragraph()
                p.add_run("Habit (Resistance): ").bold = True
                p.add_run(", ".join(fa.habit_factors[:3]))
            
            if fa.anxiety_factors:
                p = doc.add_paragraph()
                p.add_run("Anxiety Factors: ").bold = True
                p.add_run(", ".join(fa.anxiety_factors[:3]))
            
            if fa.recommended_focus:
                p = doc.add_paragraph()
                p.add_run("Recommended Focus: ").bold = True
                p.add_run(fa.recommended_focus)
    
    doc.add_page_break()


def add_journey_section(doc: Document, results: ResearchResults):
    """Add journey maps section."""
    doc.add_heading("6. Customer Journey Maps", level=1)
    
    has_journey = any(r.journey_map for r in results.icps)
    if not has_journey:
        doc.add_paragraph("Journey maps not generated.")
        doc.add_page_break()
        return
    
    for icp_result in results.icps:
        if not icp_result.journey_map:
            continue
        
        jm = icp_result.journey_map
        icp = icp_result.icp
        
        doc.add_heading(f"Journey Map: {icp.icp_name}", level=2)
        
        if jm.overall_timeline:
            p = doc.add_paragraph()
            p.add_run("Overall Timeline: ").bold = True
            p.add_run(jm.overall_timeline)
        
        stages = [
            ("Awareness", jm.awareness_stage),
            ("Consideration", jm.consideration_stage),
            ("Decision", jm.decision_stage),
            ("Onboarding", jm.onboarding_stage),
            ("Expansion", jm.expansion_stage)
        ]
        
        for stage_name, stage in stages:
            if not stage:
                continue
            
            doc.add_heading(f"Stage: {stage_name}", level=3)
            
            p = doc.add_paragraph()
            p.add_run("Objective: ").bold = True
            p.add_run(stage.objective)
            
            if stage.content_themes:
                p = doc.add_paragraph()
                p.add_run("Content Themes: ").bold = True
                p.add_run(", ".join(stage.content_themes[:5]))
            
            if stage.content_ideas:
                doc.add_paragraph("Content Ideas:", style='List Bullet')
                for idea in stage.content_ideas[:3]:
                    p = doc.add_paragraph()
                    p.add_run(f"• {idea.title}").bold = True
                    p.add_run(f" ({idea.format})")
                    if idea.hook:
                        p.add_run(f"\n  Hook: {idea.hook}")
            
            if stage.kpis:
                p = doc.add_paragraph()
                p.add_run("KPIs: ").bold = True  
                p.add_run(", ".join(stage.kpis[:5]))
    
    doc.add_page_break()


def add_cost_section(doc: Document, results: ResearchResults):
    """Add cost summary section."""
    doc.add_heading("7. Cost Summary", level=1)
    
    summary = results.cost_summary
    if not summary:
        doc.add_paragraph("Cost tracking not available.")
        return
    
    # Summary table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    metrics = [
        ("Total Cost", f"${summary.get('total_cost', 0):.4f}"),
        ("Total Requests", str(summary.get('total_requests', 0))),
        (
            "Input Tokens",
            f"{summary.get('total_tokens', {}).get('input', 0):,}"
        ),
        (
            "Output Tokens",
            f"{summary.get('total_tokens', {}).get('output', 0):,}"
        )
    ]
    
    for i, (label, value) in enumerate(metrics):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # By model breakdown
    by_model = summary.get('by_model', {})
    if by_model:
        doc.add_heading("Usage by Model", level=2)
        
        model_table = doc.add_table(
            rows=len(by_model) + 1,
            cols=4
        )
        model_table.style = 'Table Grid'
        
        # Header
        headers = ["Model", "Requests", "Tokens", "Cost"]
        for j, h in enumerate(headers):
            model_table.rows[0].cells[j].text = h
            model_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
        
        # Data
        for row_idx, (model, data) in enumerate(by_model.items(), 1):
            short_name = model.split('/')[-1]
            model_table.rows[row_idx].cells[0].text = short_name
            model_table.rows[row_idx].cells[1].text = str(data['requests'])
            model_table.rows[row_idx].cells[2].text = f"{data['total_tokens']:,}"
            model_table.rows[row_idx].cells[3].text = f"${data['cost']:.4f}"